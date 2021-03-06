
try:
	import docx
	from docx.enum.style import WD_STYLE_TYPE
	from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:

	class DocxModelReporter():
		def docx_params(self, groups=None, display_inital=False, **format):
			import docx
			from docx.enum.style import WD_STYLE_TYPE
			from docx.enum.text import WD_ALIGN_PARAGRAPH


else:
	from ..utilities import category, pmath, rename
	from ..core import LarchError, ParameterAlias


	def _append_to_document(self, other_doc):
		while not isinstance(other_doc, docx.document.Document) and hasattr(other_doc, '_parent'):
			other_doc = other_doc._parent
		if not isinstance(other_doc, docx.document.Document):
			raise TypeError('other_doc is not a docx.Document or a part thereof, it is a {}'.format(type(other_doc)))
		for element in other_doc._body._element:
			self._body._element.append(element)
		return self

	docx.document.Document.append = _append_to_document



	def document_larchstyle():
		document = docx.Document()

#		normal = document.styles['Normal']
#		normal.font.name = 'Arial'
#		normal.font.size = docx.shared.Pt(11)
#		normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
#		normal.paragraph_format.line_spacing = 1.0
#		normal.paragraph_format.widow_control = True
#
		body_text = document.styles['Body Text']

		monospaced_small = document.styles.add_style('Monospaced Small',WD_STYLE_TYPE.TABLE)
		monospaced_small.base_style = document.styles['Normal']
		monospaced_small.font.name = 'Courier New'
		monospaced_small.font.size = docx.shared.Pt(9)
		monospaced_small.paragraph_format.space_before = docx.shared.Pt(0)
		monospaced_small.paragraph_format.space_after  = docx.shared.Pt(0)
		monospaced_small.paragraph_format.line_spacing = 1.0

		table_body_text = document.styles.add_style('Table Body Text',WD_STYLE_TYPE.TABLE)
		table_body_text.base_style = document.styles['Body Text']
		table_body_text.font.name = 'Arial'
		table_body_text.font.size = docx.shared.Pt(9)
		table_body_text.paragraph_format.space_before = docx.shared.Pt(1)
		table_body_text.paragraph_format.space_after  = docx.shared.Pt(1)
		table_body_text.paragraph_format.left_indent = docx.shared.Pt(2)
		table_body_text.paragraph_format.right_indent  = docx.shared.Pt(2)
		table_body_text.paragraph_format.line_spacing = 1.0


		return document


	def docx_table(*arg, header_text=None, header_level=None,  **kwarg):
		doc = document_larchstyle()
		if header_text is not None:
			if header_level is None:
				header_level = 1
			doc.add_heading(header_text, level=header_level)
		tbl = doc.add_table(*arg, **kwarg)
		return tbl



	class DocxModelReporter():


		class DocxManager:
			"""Manages xml reporting for a :class:`Model`.	"""

			def __init__(self, model):
				self._model = model

			def _get_item(self, key, html_processor=True):
				candidate = None
				if isinstance(key,str):
					try:
						art_obj = getattr(self._model, "art_{}".format(key.casefold()))
						candidate = lambda *arg,**kwarg: art_obj(*arg,**kwarg).to_docx()
					except AttributeError:
						pass
					try:
						candidate = getattr(self._model, "docx_{}".format(key.casefold()))
					except AttributeError:
						pass
					try:
						candidate = self._model._user_defined_docx[key.casefold()]
					except (AttributeError, KeyError):
						pass
					if candidate is None:
						import warnings
						warnings.warn("docx builder for '{}' not found".format(key.casefold()))
						candidate = document_larchstyle()
						paragraph = candidate.add_paragraph(style="Body Text")
						paragraph.add_run("docx builder for '{}' not found".format(key.casefold())).bold = True
				else:
					raise TypeError("invalid item")
				return candidate

			def __getitem__(self, key):
				return self._get_item(key, True)

			def __repr__(self):
				return '<DocxManager>'

			def __str__(self):
				return repr(self)

			def __getattr__(self, key):
				if key in ('_model', ):
					return self.__dict__[key]
				return self.__getitem__(key)
		
			def __dir__(self):
				candidates = set()
				for j in dir(self._model):
					if len(j)>5 and j[:5]=='docx_':
						candidates.add(j[5:])
				try:
					self._model._user_defined_docx
				except AttributeError:
					pass
				else:
					try:
						if self._user_defined_docx == 'unpicklable local object':
							self._user_defined_docx = {}
					except AttributeError:
						self._user_defined_docx = {}
					candidates.update(i.casefold() for i in self._user_defined_docx.keys())
				return candidates
		
#			def __setattr__(self, key, val):
#				if key[0]=='_':
#					super().__setattr__(key, val)
#				else:
#					self._model.new_xhtml_section(val, key)

#			def __setitem__(self, key, val):
#				if key[0]=='_':
#					super().__setitem__(key, val)
#				else:
#					self._model.new_docx_section(val, key)

			def __call__(self, *args, filename=None, view_on_exit=False, **kwarg):
				top_doc = document_larchstyle()
				for arg in args:
					if isinstance(arg, docx.document.Document):
						top_doc.append(arg)
					elif isinstance(arg, str):
						top_doc.append(self._get_item(arg)())
					elif inspect.ismethod(arg):
						top_doc.append(arg())
				if filename is not None:
					top_doc.save(filename)
					if view_on_exit:
						try:
							os.system("open "+filename)
						except:
							pass
				return top_doc


		@property
		def docx(self):
			"""A :class:`DocxManager` interface for the model.
			
			This method creates a dicx report on the model. Call it with
			any number of string arguments to include those named report sections.
			
			All other parameters must be passed as keywords.
			
			Other Parameters
			----------------
			filename : None or str
				If None (the default) no file is generated. 
				Otherwise, this should name a file into which the html
				report will be written.  If that file already exists it will by default not 
				be overwritten, instead a new filename will be spooled off the given name.
			view_on_exit : bool
				If true, the html file will be [attempted to be] opened in Word.
				This feature may not be compatible with all platforms.
				
			Returns
			-------
			Document
			"""
			return DocxModelReporter.DocxManager(self)



		def docx_params_v0(self, groups=None, display_inital=False, **format):

			# keys fix
			existing_format_keys = list(format.keys())
			for key in existing_format_keys:
				if key.upper()!=key: format[key.upper()] = format[key]
			if 'PARAM' not in format: format['PARAM'] = '< 12.4g'
			if 'TSTAT' not in format: format['TSTAT'] = ' 0.2f'

			number_of_columns = 5
			if display_inital:
				number_of_columns += 1


			if groups is None and hasattr(self, 'parameter_groups'):
				groups = self.parameter_groups

			table = docx_table(rows=1, cols=number_of_columns, style='Table Body Text',
							   header_text="Model Parameter Estimates", header_level=2)

			def append_simple_row(name, initial_value, value, std_err, tstat, nullvalue, holdfast):
				row_cells = table.add_row().cells
				i = 0
				row_cells[i].text = name
				i += 1
				if display_inital:
					row_cells[i].text = "{:{PARAM}}".format(initial_value, **format     )
					i += 1
				row_cells[i].text = "{:{PARAM}}".format(value , **format)
				i += 1
				if holdfast:
					row_cells[i].text = "fixed value"
					row_cells[i].merge(row_cells[i+1])
					i += 2
				else:
					row_cells[i].text = "{:.3g}".format(std_err   , **format)
					i += 1
					row_cells[i].text = "{:{TSTAT}}".format(tstat , **format  )
					i += 1
				row_cells[i].text = "{:.1f}".format(nullvalue , **format)

			def append_derivative_row(name, initial_value, value, refers_to, multiplier):
				row_cells = table.add_row().cells
				i = 0
				row_cells[i].text = name
				i += 1
				if display_inital:
					row_cells[i].text = "{:{PARAM}}".format(initial_value, **format     )
					i += 1
				row_cells[i].text = "{:{PARAM}}".format(value , **format)
				i += 1
				row_cells[i].text = "= {} * {}".format(refers_to,multiplier)
				row_cells[i].merge(row_cells[i+2])
				i += 3

			hdr_cells = table.rows[0].cells
			i = 0
			hdr_cells[i].text = 'Parameter'
			i += 1
			if display_inital:
				hdr_cells[i].text = 'Initial Value'
				i += 1
			hdr_cells[i].text = 'Estimated Value'
			i += 1
			hdr_cells[i].text = 'Std Error'
			i += 1
			hdr_cells[i].text = 't-Stat'
			i += 1
			hdr_cells[i].text = 'Null Value'
			i += 1
			for cell in hdr_cells:
				cell.paragraphs[0].runs[0].bold = True


			if groups is None:
				for par in self.parameter_names():
					append_simple_row(
						par,
						self.parameter(par).initial_value,
						self.parameter(par).value,
						self.parameter(par).std_err,
						self.parameter(par).t_stat,
						self.parameter(par).null_value,
						self.parameter(par).holdfast
					)

			else:
				
				## USING GROUPS
				listed_parameters = set([p for p in groups if not isinstance(p,category)])
				for p in groups:
					if isinstance(p,category):
						listed_parameters.update( p.complete_members() )
				unlisted_parameters = (set(self.parameter_names()) | set(self.alias_names())) - listed_parameters


				def write_param_row(p, *, force=False):
					if p is None: return
					if force or (p in self) or (p in self.alias_names()):
						if isinstance(p,category):
							row_cells = table.add_row().cells
							row_cells[0].merge(row_cells[-1])
							row_cells[0].text = p.name
							#row_cells[0].style = "parameter_category"
							for subp in p.members:
								write_param_row(subp)
						else:
							if isinstance(p,rename):
								append_simple_row(par,
									self[p].initial_value,
									self[p].value,
									self[p].std_err,
									self[p].t_stat,
									self[p].null_value,
									self[p].holdfast
								)
							else:
								pwide = self.parameter_wide(p)
								if isinstance(pwide,ParameterAlias):
									append_derivative_row(pwide.name,
										self.metaparameter(pwide.name).initial_value,
										self.metaparameter(pwide.name).value,
										pwide.refers_to,
										pwide.multiplier
									)
								else:
									append_simple_row(pwide.name,
										pwide.initial_value,
										pwide.value,
										pwide.std_err,
										pwide.t_stat,
										pwide.null_value,
										pwide.holdfast
									)


				# end def
				for p in groups:
					write_param_row(p)
				if len(groups)>0 and len(unlisted_parameters)>0:
					write_param_row(category("Other Parameters"),force=True)
				if len(unlisted_parameters)>0:
					for p in unlisted_parameters:
						write_param_row(p)
			return table
		docx_param_v0 = docx_parameters_v0 = docx_params_v0


		def _docx_blurb_n(self, h_stepdown=2, n='', **format):
			try:
				blurb_rst = getattr(self, 'blurb'+str(n))
			except AttributeError:
				return None
			if isinstance(blurb_rst, bytes):
				blurb_rst = blurb_rst.decode()
			if not isinstance(blurb_rst, str):
				raise TypeError('blurb must be reStructuredText as str ot bytes')
			from ..util.rst_to_docx import render_docx
			return render_docx(blurb_rst, h_stepdown=h_stepdown)

		def docx_blurb(self, h_stepdown=2, **format):
			return self._docx_blurb_n(h_stepdown=h_stepdown, **format)

		def docx_blurb1(self, h_stepdown=2, **format):
			return self._docx_blurb_n(h_stepdown=h_stepdown, n=1, **format)

		def docx_blurb2(self, h_stepdown=2, **format):
			return self._docx_blurb_n(h_stepdown=h_stepdown, n=2, **format)

		def docx_blurb3(self, h_stepdown=2, **format):
			return self._docx_blurb_n(h_stepdown=h_stepdown, n=3, **format)

		def docx_blurb4(self, h_stepdown=2, **format):
			return self._docx_blurb_n(h_stepdown=h_stepdown, n=4, **format)

		def docx_blurb5(self, h_stepdown=2, **format):
			return self._docx_blurb_n(h_stepdown=h_stepdown, n=5, **format)

		def docx_blurb6(self, h_stepdown=2, **format):
			return self._docx_blurb_n(h_stepdown=h_stepdown, n=6, **format)

		def docx_blurb7(self, h_stepdown=2, **format):
			return self._docx_blurb_n(h_stepdown=h_stepdown, n=7, **format)

		def docx_blurb8(self, h_stepdown=2, **format):
			return self._docx_blurb_n(h_stepdown=h_stepdown, n=8, **format)

		def docx_blurb9(self, h_stepdown=2, **format):
			return self._docx_blurb_n(h_stepdown=h_stepdown, n=9, **format)


